from fastapi import HTTPException
from langchain_text_splitters import RecursiveCharacterTextSplitter
import io
import csv
import json
import docx

def extract_text(file_bytes: bytes, fileName: str) -> str:
    ext = fileName.lower().split('.')[-1]
    text = ""
    try:
        if ext == 'pdf':
            import pdfplumber
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if not table or len(table) < 2: 
                            continue
                        text_parts.append(f"\n Bảng {table_idx + 1} - Trang {page_num + 1} ")
                        
                        headers = [str(h).replace('\n', ' ').strip() if h else f"Col{j}" for j, h in enumerate(table[0])]
                        
                        for idx, row in enumerate(table[1:]):
                            row_items = []
                            for h, val in zip(headers, row):
                                if val:
                                    clean_val = str(val).replace('\n', ' ').strip()
                                    row_items.append(f"{h}: {clean_val}")
                            if row_items:
                                text_parts.append(f"[Dòng {idx+1}] " + " | ".join(row_items))
            text = "\n".join(text_parts)
        elif ext in ['doc', 'docx']:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            for i, table in enumerate(doc.tables):
                text_parts.append(f" Bảng {i+1} ")
                headers = []
                for idx, row in enumerate(table.rows):
                    cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    if idx == 0:
                        headers = [c if c else f"Column{j}" for j, c in enumerate(cells)]
                    else:
                        row_items = []
                        for h, val in zip(headers, cells):
                            if val:
                                row_items.append(f"{h}: {val}")
                        if row_items:
                            text_parts.append(f"[Dòng {idx}] " + " | ".join(row_items))
            text = "\n".join(text_parts)
        elif ext in ['csv', 'tsv']:
            if file_bytes.startswith(b'\xff\xfe') or file_bytes.startswith(b'\xfe\xff'):
                content = file_bytes.decode("utf-16", errors='ignore')
            else:
                try:
                    content = file_bytes.decode("utf-8-sig")
                except UnicodeDecodeError:
                    content = file_bytes.decode("windows-1258", errors='ignore')
            
            content = content.replace('\x00', '')
            
            # Remove empty rows or rows that only contain commas/semicolons (Excel garbage)
            lines = content.split('\n')
            valid_lines = [line for line in lines if line.replace(',', '').replace(';', '').strip()]
            content = '\n'.join(valid_lines)
            
            first_line = content.split('\n')[0] if content else ""
            if ext == 'tsv':
                delimiter = '\t'
            elif ';' in first_line and first_line.count(';') > first_line.count(','):
                delimiter = ';'
            else:
                delimiter = ','
                
            reader = csv.DictReader(io.StringIO(content), delimiter=delimiter)
            row_texts = []
            for i, row in enumerate(reader):
                row_items = [f"{k}: {v}" for k, v in row.items() if k and v and str(v).strip()]
                if row_items:
                    row_texts.append(f"[Dòng {i+1}] " + " | ".join(row_items))
            text = "\n".join(row_texts)
        elif ext == 'json':
            content = file_bytes.decode("utf-8", errors='ignore')
            try:
                text = json.dumps(json.loads(content), indent=2, ensure_ascii=False)
            except:
                text = content
        elif ext == 'xlsx':
            from openpyxl import load_workbook
            wb = load_workbook(filename=io.BytesIO(file_bytes), read_only=True, data_only=True)
            text_parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f" Bảng: {sheet_name} ")
                rows = sheet.iter_rows(values_only=True)
                try:
                    headers = next(rows)
                except StopIteration:
                    continue
                if not headers:
                    continue
                
                headers = [str(h).strip() if h is not None else f"Column{i}" for i, h in enumerate(headers)]
                for idx, row in enumerate(rows):
                    row_items = []
                    for h, val in zip(headers, row):
                        if val is not None and str(val).strip():
                            row_items.append(f"{h}: {val}")
                    if row_items:
                        text_parts.append(f"[Dòng {idx+2}] " + " | ".join(row_items))
            wb.close()
            text = "\n".join(text_parts)
        elif ext in ['txt', 'md']:
            text = file_bytes.decode("utf-8", errors='ignore')
        else:
            text = file_bytes.decode("utf-8", errors='ignore')
        if not text.strip():
            raise ValueError("Không có dữ liệu phù hợp để đọc")
    except Exception as e:
        print(f"Error extracting {fileName}: {e}")
        raise HTTPException(status_code=422, detail=f"Không thể đọc nội dung file {fileName}: {str(e)}")
    return text

def chunk_text(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size = 1500,
        chunk_overlap = 200,
        separators=["\n\n", "\n Bảng", "\n[Dòng", "\n", ".", " "]
    )
    return splitter.split_text(text)
